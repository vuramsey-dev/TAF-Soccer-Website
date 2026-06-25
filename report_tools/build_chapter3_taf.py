from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_taf_report import (
    ROOT,
    add_bullet,
    add_heading,
    add_number,
    add_para,
    add_table,
    set_run_font,
    setup_document,
)


OUT_DOCX = ROOT / "Chuong_3_TAF_Soccer.docx"


def add_command(doc, command):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 4
    p.paragraph_format.left_indent = 457200
    run = p.add_run(command)
    set_run_font(run, name="Consolas", size=11, bold=True)
    return p


def build_chapter3():
    doc = Document()
    setup_document(doc)

    add_heading(doc, "CHƯƠNG 3. CÀI ĐẶT, ĐÁNH GIÁ VÀ KẾT LUẬN", 1)
    add_para(
        doc,
        "Website bán giày bóng đá TAF Soccer được xây dựng theo mô hình client - server. Phía người dùng sử dụng trình duyệt web để truy cập các trang như đăng nhập, đăng ký, trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng và tài khoản cá nhân. Phía máy chủ sử dụng Node.js kết hợp ExpressJS để xử lý các yêu cầu API, kết nối MongoDB và trả dữ liệu về frontend.",
    )
    add_para(
        doc,
        "Hệ thống gồm ba khu vực chính: khu vực khách hàng, khu vực nhân viên và khu vực admin. Khách hàng có thể xem, tìm kiếm, chọn size, thêm giỏ hàng, đặt hàng và xem lịch sử đơn. Nhân viên quản lý sản phẩm, đơn hàng và thống kê doanh số. Admin quản lý tài khoản người dùng, phân quyền và thương hiệu sản phẩm.",
    )

    add_heading(doc, "3.1. Cài đặt, yêu cầu trang thiết bị và phần mềm hệ thống", 2)
    add_heading(doc, "3.1.1. Yêu cầu trang thiết bị", 3)
    add_para(
        doc,
        "Để cài đặt và chạy thử website TAF Soccer trong môi trường học tập, người dùng cần chuẩn bị máy tính cá nhân hoặc laptop có cấu hình cơ bản. Vì hệ thống chạy cục bộ với Node.js, MongoDB và trình duyệt web nên yêu cầu phần cứng không quá cao.",
    )
    add_table(
        doc,
        ["Thành phần", "Yêu cầu tối thiểu", "Khuyến nghị"],
        [
            ["CPU", "Intel Core i3 hoặc tương đương", "Intel Core i5 trở lên"],
            ["RAM", "4GB", "8GB trở lên"],
            ["Ổ cứng", "Còn trống tối thiểu 5GB", "SSD còn trống 10GB trở lên"],
            ["Mạng", "Có kết nối Internet khi cài thư viện", "Kết nối ổn định để kiểm thử và tra cứu tài liệu"],
            ["Màn hình", "Độ phân giải HD", "Full HD để kiểm thử giao diện thuận tiện hơn"],
            ["Thiết bị kiểm thử", "Máy tính cá nhân", "Thêm điện thoại hoặc chế độ responsive của trình duyệt"],
        ],
        widths_dxa=[1900, 3500, 3960],
        font_size=10.8,
    )
    add_para(
        doc,
        "Với cấu hình trên, hệ thống có thể chạy tốt trong phạm vi học tập, trình bày sản phẩm và kiểm thử chức năng cơ bản. Khi triển khai thực tế cho nhiều người dùng truy cập đồng thời, cần sử dụng máy chủ riêng hoặc nền tảng cloud có tài nguyên cao hơn.",
    )

    add_heading(doc, "3.1.2. Yêu cầu phần mềm", 3)
    add_para(doc, "Các phần mềm và công cụ cần thiết để cài đặt, chạy thử và phát triển hệ thống gồm:")
    add_table(
        doc,
        ["Phần mềm/Công cụ", "Vai trò trong hệ thống"],
        [
            ["Windows 10/11", "Hệ điều hành dùng để phát triển và chạy thử project."],
            ["Visual Studio Code", "Môi trường chỉnh sửa mã nguồn HTML, CSS, JavaScript và Node.js."],
            ["Node.js và npm", "Chạy backend ExpressJS và cài đặt thư viện trong package.json."],
            ["MongoDB", "Lưu trữ dữ liệu sản phẩm, người dùng, giỏ hàng, đơn hàng, thương hiệu và địa chỉ."],
            ["Google Chrome hoặc Microsoft Edge", "Kiểm thử giao diện khách hàng, nhân viên và admin."],
            ["Git hoặc công cụ nén/giải nén", "Quản lý, sao chép hoặc nộp mã nguồn project."],
        ],
        widths_dxa=[2600, 6760],
    )
    add_para(doc, "Trong project, các thư viện backend chính được khai báo trong file package.json:")
    add_table(
        doc,
        ["Thư viện", "Chức năng"],
        [
            ["express", "Xây dựng server backend và định nghĩa các REST API."],
            ["mongoose", "Kết nối MongoDB, định nghĩa schema/model và thao tác dữ liệu."],
            ["cors", "Cho phép frontend gọi API từ trình duyệt trong quá trình chạy thử."],
            ["nodemon", "Tự động khởi động lại server khi mã nguồn backend thay đổi trong chế độ phát triển."],
        ],
        widths_dxa=[2100, 7260],
    )
    add_para(
        doc,
        "Việc sử dụng Node.js, ExpressJS và MongoDB giúp hệ thống có cấu trúc rõ ràng, phù hợp với mô hình website bán hàng hiện đại. Dữ liệu dạng document của MongoDB thuận tiện cho việc lưu trữ sản phẩm có nhiều size, giỏ hàng, đơn hàng và địa chỉ nhận hàng.",
    )

    add_heading(doc, "3.1.3. Cấu trúc thư mục cài đặt", 3)
    add_para(
        doc,
        "Source code của website được tổ chức theo hướng tách biệt giữa backend, frontend, tài nguyên ảnh, file CSS và file JavaScript theo từng trang. Cách tổ chức này giúp việc lập trình, kiểm thử và bảo trì thuận tiện hơn.",
    )
    add_table(
        doc,
        ["Thư mục/Tệp", "Chức năng"],
        [
            ["package.json", "Khai báo thông tin project, script start/dev và các thư viện cần dùng."],
            ["Backend/server.js", "Khởi tạo Express, kết nối MongoDB taf_soccer, khai báo route API và phục vụ thư mục Frontend."],
            ["Backend/models", "Chứa các model User, Product, Cart, Order, Brand và Address."],
            ["Backend/routes", "Chứa các route xử lý auth, products, cart, orders, statistics, admin và addresses."],
            ["Frontend/user", "Các trang dành cho khách hàng như login, register, homepage, all-products, product-detail, cart, account."],
            ["Frontend/staff", "Các trang dành cho nhân viên như dashboard, products, orders và statistics."],
            ["Frontend/admin", "Các trang admin quản lý người dùng và thương hiệu."],
            ["Frontend/JS", "Các file JavaScript xử lý gọi API, sự kiện giao diện, giỏ hàng, sản phẩm, đơn hàng và thống kê."],
            ["Frontend/CSS", "Các file CSS định dạng giao diện theo từng nhóm trang."],
            ["Frontend/img", "Lưu logo, ảnh minh họa và ảnh mặc định dùng trong giao diện."],
        ],
        widths_dxa=[2600, 6760],
        font_size=11,
    )
    add_para(
        doc,
        "Trong đó, Backend đóng vai trò xử lý nghiệp vụ và giao tiếp với cơ sở dữ liệu. Frontend chịu trách nhiệm hiển thị giao diện và gửi yêu cầu đến API thông qua hàm fetch của JavaScript.",
    )

    add_heading(doc, "3.1.4. Các bước cài đặt và chạy hệ thống", 3)
    add_para(doc, "Quy trình cài đặt hệ thống được thực hiện theo các bước sau:")
    steps = [
        ("Bước 1: Cài đặt Node.js và npm", ["node -v", "npm -v"], "Nếu hai lệnh trên hiển thị phiên bản, môi trường Node.js đã được cài đặt thành công."),
        ("Bước 2: Cài đặt và khởi động MongoDB", [], "Người dùng cần đảm bảo MongoDB local đang hoạt động. Server hiện kết nối đến địa chỉ mongodb://127.0.0.1:27017/taf_soccer."),
        ("Bước 3: Mở project bằng Visual Studio Code", [], "Mở đúng thư mục gốc TAF_Soccer_fixed, nơi chứa file package.json. Đây là thư mục dùng để chạy các lệnh npm."),
        ("Bước 4: Cài đặt thư viện", ["npm install"], "Lệnh này cài các thư viện express, mongoose, cors và nodemon theo khai báo trong package.json."),
        ("Bước 5: Chạy server backend", ["npm start", "npm run dev"], "Có thể dùng npm start để chạy bằng Node.js hoặc npm run dev để chạy với nodemon trong quá trình phát triển."),
        ("Bước 6: Truy cập website", ["http://localhost:3000", "http://localhost:3000/Frontend/user/login.html"], "Khi mở địa chỉ gốc, server tự chuyển người dùng đến trang đăng nhập của khu vực khách hàng."),
    ]
    for title, commands, desc in steps:
        add_para(doc, title, bold=True, after=2)
        for command in commands:
            add_command(doc, command)
        add_para(doc, desc)
    add_para(
        doc,
        "Sau khi server chạy thành công, người dùng có thể đăng ký tài khoản khách hàng, đăng nhập, xem sản phẩm và kiểm thử các chức năng chính. Với tài khoản staff hoặc admin, hệ thống sẽ chuyển đến các trang quản trị tương ứng theo vai trò.",
    )

    add_heading(doc, "3.1.5. Yêu cầu dữ liệu ban đầu", 3)
    add_para(
        doc,
        "Để website hoạt động trực quan khi chạy thử, hệ thống cần có dữ liệu ban đầu gồm tài khoản người dùng, tài khoản nhân viên, tài khoản admin, danh sách thương hiệu và danh sách sản phẩm giày bóng đá.",
    )
    add_table(
        doc,
        ["Loại dữ liệu", "Nội dung cần có", "Mục đích"],
        [
            ["User", "Tài khoản khách hàng, nhân viên và admin với role tương ứng user, staff, admin.", "Kiểm thử đăng nhập, phân quyền và luồng mua hàng."],
            ["Brand", "Các thương hiệu giày như Nike, Adidas, Puma, Mizuno hoặc thương hiệu do cửa hàng quản lý.", "Phục vụ lọc, quản lý và đồng bộ thương hiệu sản phẩm."],
            ["Product", "Tên giày, thương hiệu, giá, ảnh, mô tả và sizeStock theo từng size.", "Hiển thị danh sách sản phẩm, chi tiết sản phẩm và tồn kho."],
            ["Cart", "Được tạo khi người dùng thao tác với giỏ hàng.", "Lưu sản phẩm đã chọn theo userId, productId, size và quantity."],
            ["Order", "Được tạo sau khi khách hàng đặt hàng.", "Lưu thông tin người nhận, sản phẩm, tổng tiền, trạng thái đơn và thanh toán."],
            ["Address", "Địa chỉ nhận hàng đã lưu của khách hàng.", "Cho phép chọn lại địa chỉ khi đặt hàng lần sau."],
        ],
        widths_dxa=[1600, 4500, 3260],
        font_size=10.5,
    )
    add_para(
        doc,
        "Nếu cơ sở dữ liệu MongoDB chưa có dữ liệu, admin hoặc người phát triển có thể tạo dữ liệu thủ công bằng trang quản trị, MongoDB Compass hoặc script seed riêng. Sau khi có dữ liệu mẫu, quá trình trình bày sản phẩm sẽ thuận tiện hơn.",
    )

    add_heading(doc, "3.1.6. Tài khoản chạy thử hệ thống", 3)
    add_para(
        doc,
        "Để kiểm tra các quyền khác nhau trong website, hệ thống cần có các tài khoản chạy thử tương ứng với ba vai trò chính. Các tài khoản dưới đây là dữ liệu demo đề xuất, có thể tạo thủ công trong MongoDB hoặc qua chức năng quản lý người dùng của admin.",
    )
    add_table(
        doc,
        ["Vai trò", "Email demo đề xuất", "Mật khẩu demo", "Mục đích sử dụng"],
        [
            ["Khách hàng", "user@tafsoccer.local", "123456", "Kiểm thử xem sản phẩm, giỏ hàng, đặt hàng và lịch sử đơn."],
            ["Nhân viên", "staff@tafsoccer.local", "123456", "Kiểm thử quản lý sản phẩm, đơn hàng và thống kê."],
            ["Admin", "admin@tafsoccer.local", "123456", "Kiểm thử quản lý người dùng, phân quyền và thương hiệu."],
        ],
        widths_dxa=[1500, 2900, 1700, 3260],
        font_size=10.5,
    )
    add_para(
        doc,
        "Lưu ý: các tài khoản trên chỉ dùng cho mục đích trình bày và kiểm thử trong môi trường học tập. Khi triển khai thật, không nên sử dụng mật khẩu đơn giản hoặc công khai thông tin đăng nhập quản trị.",
    )

    add_heading(doc, "3.2. Lập trình, tích hợp hệ thống", 2)
    add_heading(doc, "3.2.1. Lập trình phía giao diện người dùng", 3)
    add_para(
        doc,
        "Giao diện người dùng được xây dựng bằng HTML, CSS và JavaScript thuần. Khu vực khách hàng gồm các trang đăng nhập, đăng ký, trang chủ, danh sách sản phẩm, sản phẩm theo thương hiệu, chi tiết sản phẩm, giỏ hàng và tài khoản cá nhân.",
    )
    add_para(
        doc,
        "Các file CSS trong thư mục Frontend/CSS định dạng bố cục, màu sắc, menu, sản phẩm, form, bảng, giỏ hàng và trang quản trị. Các file JavaScript trong Frontend/JS xử lý sự kiện người dùng, gọi API bằng fetch, đọc/ghi localStorage và cập nhật giao diện theo dữ liệu trả về từ backend.",
    )
    add_para(doc, "Các chức năng chính phía frontend gồm:")
    for item in [
        "Hiển thị danh sách sản phẩm và sản phẩm theo từng thương hiệu.",
        "Tìm kiếm, lọc và sắp xếp sản phẩm theo tên, thương hiệu, giá hoặc trạng thái còn hàng.",
        "Lưu selectedProductId vào localStorage khi người dùng chọn xem chi tiết sản phẩm.",
        "Lưu thông tin đăng nhập người dùng vào localStorage để phân luồng theo role.",
        "Hiển thị giỏ hàng, tính tổng tiền, phí vận chuyển và số lượng sản phẩm đã chọn.",
        "Hiển thị dashboard, bảng quản lý sản phẩm, đơn hàng, người dùng, thương hiệu và thống kê.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2.2. Lập trình chức năng xem chi tiết sản phẩm và giỏ hàng", 3)
    add_para(
        doc,
        "Khi người dùng chọn một sản phẩm, frontend chuyển sang trang product-detail.html và lấy mã sản phẩm từ localStorage. File product-detail.js gọi API /api/products/:id để lấy dữ liệu chi tiết, sau đó hiển thị tên sản phẩm, thương hiệu, giá bán, ảnh, mô tả, danh sách size còn hàng và số lượng tồn.",
    )
    add_para(
        doc,
        "Chức năng giỏ hàng được xử lý trong cart.js và cartRoutes.js. Mỗi sản phẩm trong giỏ lưu productId, name, brand, price, image, size và quantity. Việc lưu size trong từng dòng giỏ hàng phù hợp với đặc thù giày bóng đá vì cùng một mẫu giày có thể có nhiều size khác nhau.",
    )
    add_para(doc, "Các chức năng chính của giỏ hàng gồm:")
    for item in [
        "Thêm sản phẩm vào giỏ theo userId, productId, size và quantity.",
        "Nếu sản phẩm cùng size đã tồn tại, hệ thống tăng số lượng thay vì tạo dòng mới.",
        "Cập nhật số lượng sản phẩm trong giỏ hàng.",
        "Xóa một sản phẩm, xóa nhiều sản phẩm đã chọn hoặc làm trống toàn bộ giỏ hàng.",
        "Tính tạm tính, phí vận chuyển và tổng tiền của các sản phẩm được chọn.",
        "Cập nhật biểu tượng số lượng sản phẩm trong giỏ hàng trên menu.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2.3. Lập trình chức năng đặt hàng", 3)
    add_para(
        doc,
        "Sau khi chọn sản phẩm trong giỏ hàng, khách hàng có thể tiến hành đặt hàng. Hệ thống hiển thị form nhận hàng gồm tên người nhận, số điện thoại, tỉnh/thành phố, xã/phường và địa chỉ chi tiết. Người dùng cũng có thể lưu địa chỉ để dùng lại trong các lần đặt hàng sau.",
    )
    add_para(
        doc,
        "Khi khách hàng xác nhận đặt hàng, frontend gửi dữ liệu đến API /api/orders. Backend kiểm tra userId, thông tin người nhận, danh sách sản phẩm và tổng tiền. Nếu dữ liệu hợp lệ, hệ thống tạo bản ghi Order với trạng thái mặc định “Chờ xác nhận” và trạng thái thanh toán “Chưa thanh toán”. Sau đó, các sản phẩm đã đặt được xóa khỏi giỏ hàng.",
    )
    add_table(
        doc,
        ["Bước", "Nội dung xử lý", "Kết quả"],
        [
            ["1", "Khách hàng chọn sản phẩm cần thanh toán trong giỏ hàng.", "Danh sách sản phẩm được đưa vào selectedCheckoutProducts."],
            ["2", "Khách hàng nhập hoặc chọn địa chỉ nhận hàng đã lưu.", "Thông tin người nhận được kiểm tra đầy đủ."],
            ["3", "Frontend gửi POST /api/orders.", "Dữ liệu đơn hàng được chuyển tới backend."],
            ["4", "Backend lưu Order vào MongoDB.", "Đơn hàng mới được tạo với trạng thái Chờ xác nhận."],
            ["5", "Backend cập nhật Cart.", "Các sản phẩm đã đặt bị xóa khỏi giỏ hàng."],
            ["6", "Frontend hiển thị thông báo đặt hàng thành công.", "Khách hàng có thể xem đơn hàng trong trang tài khoản."],
        ],
        widths_dxa=[800, 5200, 3360],
        font_size=10.7,
    )

    add_heading(doc, "3.2.4. Lập trình trang thương hiệu và danh sách sản phẩm", 3)
    add_para(
        doc,
        "Website có các trang sản phẩm theo thương hiệu như adidas.html, nike.html, puma.html và mizuno.html. File brand-products.js gọi API lấy danh sách sản phẩm, lọc theo thương hiệu tương ứng và hiển thị dưới dạng lưới sản phẩm.",
    )
    add_para(
        doc,
        "Trang all-products.html hiển thị toàn bộ sản phẩm. File all-products.js hỗ trợ thống kê số sản phẩm, lọc theo size còn hàng, tìm kiếm theo từ khóa và sắp xếp sản phẩm. Người dùng có thể chuyển sang trang chi tiết hoặc thêm sản phẩm vào giỏ hàng trực tiếp từ danh sách.",
    )

    add_heading(doc, "3.2.5. Lập trình trang quản trị nhân viên và admin", 3)
    add_para(
        doc,
        "Khu vực nhân viên được xây dựng trong thư mục Frontend/staff, gồm dashboard, quản lý sản phẩm, quản lý đơn hàng và thống kê. Khu vực admin được xây dựng trong thư mục Frontend/admin, gồm quản lý người dùng và quản lý thương hiệu.",
    )
    add_para(doc, "Các chức năng quản trị chính gồm:")
    for item in [
        "Nhân viên thêm, sửa, xóa sản phẩm; quản lý tồn kho theo từng size.",
        "Nhân viên tải ảnh sản phẩm dạng base64 và xem ảnh xem trước trước khi lưu.",
        "Nhân viên theo dõi danh sách đơn hàng, xem chi tiết đơn và cập nhật trạng thái theo luồng nghiệp vụ.",
        "Nhân viên đánh dấu thanh toán sau khi đơn hàng đã giao.",
        "Admin thêm, sửa, xóa tài khoản và phân quyền user, staff, admin.",
        "Admin quản lý thương hiệu, đồng bộ tên thương hiệu sang sản phẩm, giỏ hàng và đơn hàng khi đổi tên.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "Hệ thống kiểm tra vai trò người dùng ở frontend. Tài khoản staff được chuyển về khu vực nhân viên, tài khoản admin được chuyển về khu vực admin, còn tài khoản user được sử dụng khu vực khách hàng.",
    )

    add_heading(doc, "3.2.6. Lập trình thống kê doanh số", 3)
    add_para(
        doc,
        "Chức năng thống kê được xây dựng bằng statisticRoutes.js ở backend và staff-statistics.js ở frontend. API /api/statistics/sales tổng hợp dữ liệu từ Product và Order để trả về các chỉ số phục vụ quản lý.",
    )
    add_table(
        doc,
        ["Chỉ số thống kê", "Ý nghĩa"],
        [
            ["totalProducts", "Tổng số sản phẩm đang có trong hệ thống."],
            ["totalOrders", "Tổng số đơn hàng trong khoảng thời gian lọc."],
            ["totalCompletedOrders", "Số đơn hàng đã hoàn thành hoặc đã thanh toán."],
            ["totalRevenue", "Tổng doanh thu từ các đơn đủ điều kiện tính doanh thu."],
            ["totalSoldProducts", "Tổng số lượng sản phẩm đã bán."],
            ["revenueByMonth", "Doanh thu và số đơn hoàn thành theo từng tháng."],
            ["topProducts", "Danh sách sản phẩm bán chạy theo số lượng."],
            ["orderStatusList", "Số lượng đơn hàng theo từng trạng thái."],
        ],
        widths_dxa=[2600, 6760],
    )
    add_para(
        doc,
        "Frontend hiển thị các chỉ số thống kê, bảng doanh thu theo tháng, bảng sản phẩm bán chạy, bảng trạng thái đơn hàng và biểu đồ cột đơn giản. Ngoài ra, trang thống kê còn có chức năng xuất dữ liệu ra file CSV để mở bằng Excel.",
    )

    add_heading(doc, "3.2.7. Lập trình backend bằng Node.js và Express", 3)
    add_para(
        doc,
        "Backend của hệ thống được xây dựng trong thư mục Backend. File server.js khởi tạo Express, bật CORS, cấu hình đọc JSON với giới hạn 10MB, phục vụ thư mục Frontend dưới đường dẫn /Frontend và kết nối MongoDB bằng Mongoose.",
    )
    add_table(
        doc,
        ["Route", "Chức năng chính"],
        [
            ["authRoutes", "Đăng ký, đăng nhập và đổi mật khẩu."],
            ["productRoutes", "Lấy danh sách, thêm, sửa, xóa, tìm kiếm và xem chi tiết sản phẩm."],
            ["cartRoutes", "Lấy giỏ hàng, thêm sản phẩm, cập nhật số lượng, xóa sản phẩm và làm trống giỏ hàng."],
            ["orderRoutes", "Tạo đơn hàng, lấy lịch sử đơn, lấy danh sách đơn, cập nhật trạng thái và thanh toán."],
            ["adminRoutes", "Quản lý người dùng, phân quyền và quản lý thương hiệu."],
            ["statisticRoutes", "Thống kê doanh số, sản phẩm bán chạy và trạng thái đơn hàng."],
            ["addressRoutes", "Lưu, lấy và xóa địa chỉ nhận hàng."],
        ],
        widths_dxa=[2200, 7160],
    )

    add_heading(doc, "3.2.8. Xây dựng cơ chế xác thực và phân quyền", 3)
    add_para(
        doc,
        "Hệ thống có chức năng đăng ký, đăng nhập và đổi mật khẩu trong authRoutes.js. Khi đăng ký, người dùng nhập họ tên, email, mật khẩu và số điện thoại. Backend kiểm tra dữ liệu bắt buộc, kiểm tra độ dài mật khẩu và kiểm tra email đã tồn tại hay chưa.",
    )
    add_para(
        doc,
        "Khi đăng nhập thành công, backend trả về thông tin id, name, email, phone và role. Frontend lưu thông tin này vào localStorage, sau đó điều hướng người dùng theo vai trò: user vào trang khách hàng, staff vào trang nhân viên và admin vào trang quản trị.",
    )
    add_para(
        doc,
        "Cơ chế phân quyền hiện tại phù hợp với phạm vi bài tập lớn và chạy thử cục bộ. Tuy nhiên, khi triển khai thực tế cần bổ sung JWT/session, mã hóa mật khẩu bằng bcrypt, bảo vệ API phía backend và sử dụng HTTPS để nâng cao an toàn dữ liệu.",
    )

    add_heading(doc, "3.2.9. Tích hợp cơ sở dữ liệu MongoDB", 3)
    add_para(
        doc,
        "MongoDB được sử dụng để lưu trữ dữ liệu của hệ thống. Mỗi nhóm dữ liệu được định nghĩa thành một Mongoose model riêng, giúp backend thao tác dữ liệu rõ ràng và dễ mở rộng.",
    )
    add_table(
        doc,
        ["Model", "Trường dữ liệu chính", "Chức năng"],
        [
            ["User", "name, email, password, phone, role, timestamps", "Lưu tài khoản khách hàng, nhân viên và admin."],
            ["Product", "name, brand, price, sizeStock, quantity, image, description", "Lưu thông tin giày bóng đá và tồn kho theo size."],
            ["Cart", "userId, products(productId, name, brand, price, image, size, quantity)", "Lưu giỏ hàng của từng người dùng."],
            ["Order", "userId, customerName, phone, address, products, totalPrice, status, paymentStatus, paidAt, createdAt", "Lưu đơn hàng và trạng thái xử lý."],
            ["Brand", "name, description, timestamps", "Lưu thương hiệu và phục vụ đồng bộ với sản phẩm."],
            ["Address", "userId, name, phone, province, ward, detailAddress, address, createdAt", "Lưu địa chỉ nhận hàng của khách hàng."],
        ],
        widths_dxa=[1300, 4700, 3360],
        font_size=10,
    )
    add_para(
        doc,
        "Trong model Product, trường sizeStock là mảng gồm size và quantity, cho phép quản lý số lượng tồn theo từng size. Khi thêm hoặc cập nhật sản phẩm, backend chuẩn hóa sizeStock và tính tổng quantity để hiển thị nhanh trên giao diện quản lý.",
    )
    add_para(
        doc,
        "Trong model Order, trạng thái đơn hàng gồm: Chờ xác nhận, Đã xác nhận, Đang giao, Đã giao và Hoàn thành. Trạng thái thanh toán gồm: Chưa thanh toán và Đã thanh toán. Backend chỉ cho phép hoàn thành đơn sau khi đơn đã được đánh dấu thanh toán.",
    )

    add_heading(doc, "3.2.10. Tích hợp các API chính", 3)
    add_para(doc, "Các API chính của hệ thống được tổ chức theo từng nhóm nghiệp vụ:")
    add_table(
        doc,
        ["Nhóm API", "Phương thức/Đường dẫn", "Chức năng"],
        [
            ["Auth", "POST /api/auth/register", "Đăng ký tài khoản khách hàng."],
            ["Auth", "POST /api/auth/login", "Đăng nhập hệ thống."],
            ["Auth", "PUT /api/auth/change-password", "Đổi mật khẩu tài khoản."],
            ["Products", "GET /api/products", "Lấy danh sách sản phẩm."],
            ["Products", "POST /api/products", "Thêm sản phẩm mới."],
            ["Products", "PUT /api/products/:id", "Cập nhật sản phẩm."],
            ["Products", "DELETE /api/products/:id", "Xóa sản phẩm."],
            ["Products", "GET /api/products/search/:keyword", "Tìm kiếm sản phẩm."],
            ["Cart", "GET /api/cart/:userId", "Lấy giỏ hàng của người dùng."],
            ["Cart", "POST /api/cart/add", "Thêm sản phẩm vào giỏ hàng."],
            ["Cart", "PUT /api/cart/update", "Cập nhật số lượng trong giỏ."],
            ["Cart", "DELETE /api/cart/remove", "Xóa sản phẩm khỏi giỏ."],
            ["Orders", "POST /api/orders", "Tạo đơn hàng mới."],
            ["Orders", "GET /api/orders/user/:userId", "Lấy lịch sử đơn hàng của khách."],
            ["Orders", "PUT /api/orders/:id/status", "Cập nhật trạng thái đơn hàng."],
            ["Orders", "PUT /api/orders/:id/payment", "Đánh dấu thanh toán."],
            ["Admin", "GET/POST/PUT/DELETE /api/admin/users", "Quản lý tài khoản người dùng."],
            ["Admin", "GET/POST/PUT/DELETE /api/admin/brands", "Quản lý thương hiệu."],
            ["Statistics", "GET /api/statistics/sales", "Thống kê doanh số."],
            ["Address", "GET/POST/DELETE /api/addresses", "Quản lý địa chỉ nhận hàng."],
        ],
        widths_dxa=[1400, 3400, 4560],
        font_size=9.5,
    )
    add_para(
        doc,
        "Việc tách API theo nhóm giúp hệ thống dễ theo dõi, dễ kiểm thử và thuận tiện khi mở rộng thêm các chức năng như thanh toán online, vận chuyển hoặc đánh giá sản phẩm.",
    )

    add_heading(doc, "3.2.11. Kết quả lập trình và tích hợp", 3)
    add_para(doc, "Sau khi lập trình và tích hợp, hệ thống đạt được các kết quả sau:")
    for item in [
        "Website chạy được trên trình duyệt thông qua server Express tại cổng 3000.",
        "Trang khách hàng có đăng nhập, đăng ký, trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng và tài khoản.",
        "Khách hàng có thể chọn size, thêm sản phẩm vào giỏ, cập nhật số lượng và đặt hàng.",
        "Đơn hàng được lưu vào MongoDB và sản phẩm đã đặt được xóa khỏi giỏ hàng.",
        "Nhân viên có thể quản lý sản phẩm, tồn kho theo size, đơn hàng và thống kê.",
        "Admin có thể quản lý người dùng, phân quyền và thương hiệu sản phẩm.",
        "Dữ liệu người dùng, sản phẩm, giỏ hàng, đơn hàng, thương hiệu và địa chỉ được lưu trong MongoDB.",
        "Các chức năng chính đã bám sát yêu cầu phân tích và thiết kế của hệ thống bán giày bóng đá.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3. Thử nghiệm, đánh giá hệ thống, kết luận", 2)
    add_heading(doc, "3.3.1. Mục tiêu thử nghiệm", 3)
    add_para(
        doc,
        "Mục tiêu thử nghiệm là kiểm tra xem các chức năng chính của website TAF Soccer có hoạt động đúng theo yêu cầu đã phân tích hay không. Việc thử nghiệm tập trung vào các luồng nghiệp vụ quan trọng như đăng ký, đăng nhập, xem sản phẩm, chọn size, giỏ hàng, đặt hàng, quản lý sản phẩm, quản lý đơn hàng, quản lý người dùng, quản lý thương hiệu và thống kê.",
    )
    add_para(doc, "Các mục tiêu cụ thể gồm:")
    for item in [
        "Kiểm tra website có mở đúng trang đăng nhập và các trang khách hàng hay không.",
        "Kiểm tra chức năng đăng ký, đăng nhập và điều hướng theo role.",
        "Kiểm tra danh sách sản phẩm có được lấy từ MongoDB hay không.",
        "Kiểm tra chức năng tìm kiếm, lọc, sắp xếp và xem chi tiết sản phẩm.",
        "Kiểm tra chức năng chọn size, thêm giỏ hàng, cập nhật số lượng và xóa sản phẩm.",
        "Kiểm tra chức năng lưu địa chỉ nhận hàng và tạo đơn hàng.",
        "Kiểm tra chức năng quản lý sản phẩm, thương hiệu, người dùng và đơn hàng.",
        "Kiểm tra chức năng thống kê doanh số và xuất CSV.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3.2. Kịch bản thử nghiệm", 3)
    add_para(doc, "Bảng dưới đây trình bày một số kịch bản thử nghiệm chức năng chính của hệ thống:")
    add_table(
        doc,
        ["STT", "Chức năng", "Dữ liệu/Thao tác thử nghiệm", "Kết quả mong đợi", "Kết quả thực tế", "Đánh giá"],
        [
            ["1", "Đăng ký tài khoản", "Nhập họ tên, email, mật khẩu từ 6 ký tự", "Tài khoản được tạo với role user", "Hệ thống thông báo đăng ký thành công", "Đạt"],
            ["2", "Đăng nhập", "Nhập email và mật khẩu đúng", "Đăng nhập thành công và điều hướng theo role", "User vào trang khách hàng, staff/admin vào trang quản trị tương ứng", "Đạt"],
            ["3", "Xem danh sách sản phẩm", "Mở trang tất cả sản phẩm", "Sản phẩm được tải từ API /api/products", "Danh sách sản phẩm hiển thị trên giao diện", "Đạt"],
            ["4", "Tìm kiếm sản phẩm", "Nhập từ khóa theo tên hoặc thương hiệu", "Hiển thị sản phẩm phù hợp", "Kết quả lọc đúng theo từ khóa", "Đạt"],
            ["5", "Xem chi tiết sản phẩm", "Chọn một sản phẩm", "Hiển thị ảnh, giá, mô tả, size và tồn kho", "Trang chi tiết hiển thị đúng dữ liệu", "Đạt"],
            ["6", "Thêm vào giỏ hàng", "Chọn size và số lượng", "Sản phẩm được thêm vào giỏ theo đúng size", "Giỏ hàng cập nhật đúng", "Đạt"],
            ["7", "Cập nhật giỏ hàng", "Tăng/giảm số lượng hoặc xóa sản phẩm", "Giỏ hàng và tổng tiền thay đổi đúng", "Số lượng và tổng tiền được cập nhật", "Đạt"],
            ["8", "Đặt hàng", "Chọn sản phẩm, nhập địa chỉ nhận hàng", "Đơn hàng được tạo và sản phẩm bị xóa khỏi giỏ", "Thông báo đặt hàng thành công", "Đạt"],
            ["9", "Lịch sử đơn hàng", "Mở tab đơn hàng trong tài khoản", "Hiển thị các đơn của người dùng", "Danh sách đơn hàng được tải theo userId", "Đạt"],
            ["10", "Quản lý sản phẩm", "Nhân viên thêm/sửa/xóa sản phẩm", "Dữ liệu Product được cập nhật", "Danh sách sản phẩm thay đổi đúng", "Đạt"],
            ["11", "Quản lý đơn hàng", "Cập nhật trạng thái đơn theo luồng", "Đơn chuyển trạng thái hợp lệ", "Backend kiểm tra và cập nhật đúng", "Đạt"],
            ["12", "Đánh dấu thanh toán", "Đơn ở trạng thái Đã giao", "Thanh toán được ghi nhận", "paymentStatus chuyển sang Đã thanh toán", "Đạt"],
            ["13", "Quản lý người dùng", "Admin thêm/sửa/xóa tài khoản", "Danh sách tài khoản cập nhật", "Dữ liệu User thay đổi đúng", "Đạt"],
            ["14", "Quản lý thương hiệu", "Admin thêm/sửa/xóa thương hiệu", "Brand được cập nhật, không xóa khi còn sản phẩm", "Hệ thống kiểm tra đúng điều kiện", "Đạt"],
            ["15", "Thống kê doanh số", "Mở trang thống kê hoặc lọc ngày", "Hiển thị doanh thu, top sản phẩm, trạng thái đơn", "Số liệu và biểu đồ hiển thị", "Đạt"],
        ],
        widths_dxa=[550, 1500, 2500, 2300, 1900, 610],
        font_size=8.8,
    )

    add_heading(doc, "3.3.3. Kết quả thử nghiệm", 3)
    add_para(
        doc,
        "Qua quá trình thử nghiệm các chức năng chính, hệ thống đáp ứng được yêu cầu cơ bản của một website bán giày bóng đá trực tuyến. Các chức năng hiển thị sản phẩm, tìm kiếm, xem chi tiết, chọn size, thêm giỏ hàng và đặt hàng hoạt động ổn định trong môi trường chạy thử cục bộ.",
    )
    add_para(
        doc,
        "Chức năng đăng ký và đăng nhập xử lý đúng dữ liệu đầu vào. Khi đăng ký, hệ thống kiểm tra họ tên, email, mật khẩu và email trùng. Khi đăng nhập thành công, thông tin người dùng được lưu ở localStorage để phục vụ việc hiển thị giao diện và điều hướng theo vai trò.",
    )
    add_para(
        doc,
        "Chức năng giỏ hàng hoạt động đúng theo yêu cầu. Sản phẩm được lưu theo từng size, số lượng và người dùng. Khi khách hàng đặt hàng thành công, các sản phẩm đã đặt được xóa khỏi giỏ hàng, giúp tránh tình trạng đặt trùng sản phẩm.",
    )
    add_para(
        doc,
        "Khu vực nhân viên và admin hiển thị được các nhóm chức năng cần thiết. Nhân viên có thể quản lý sản phẩm, đơn hàng và thống kê. Admin có thể quản lý tài khoản, phân quyền và thương hiệu. Các chức năng này phù hợp với yêu cầu của một hệ thống bán hàng trực tuyến quy mô học tập.",
    )

    add_heading(doc, "3.3.4. Đánh giá ưu điểm", 3)
    add_para(doc, "Hệ thống có một số ưu điểm sau:")
    for item in [
        "Giao diện được chia rõ theo ba vai trò: khách hàng, nhân viên và admin.",
        "Luồng mua hàng rõ ràng, gồm xem sản phẩm, chọn size, thêm giỏ hàng, nhập địa chỉ và đặt hàng.",
        "Sản phẩm có quản lý tồn kho theo từng size, phù hợp với đặc thù giày bóng đá.",
        "Có chức năng lưu địa chỉ nhận hàng, giúp khách hàng đặt hàng thuận tiện hơn.",
        "Có quản lý trạng thái đơn hàng theo trình tự nghiệp vụ rõ ràng.",
        "Có thống kê doanh thu, sản phẩm bán chạy và trạng thái đơn hàng.",
        "Dữ liệu được lưu bằng MongoDB, phù hợp với ứng dụng web hiện đại.",
        "Source code được chia thành Backend, Frontend, models, routes, CSS và JS theo từng trang, giúp dễ bảo trì.",
        "Có chức năng quản lý thương hiệu và đồng bộ tên thương hiệu khi cập nhật.",
        "Có thể mở rộng thêm các chức năng thanh toán, vận chuyển và đánh giá sản phẩm trong tương lai.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3.5. Hạn chế của hệ thống", 3)
    add_para(
        doc,
        "Bên cạnh các kết quả đã đạt được, website TAF Soccer vẫn còn một số hạn chế do phạm vi đề tài là bài tập lớn và hệ thống chủ yếu được triển khai trong môi trường chạy thử cục bộ.",
    )
    for item in [
        "Hệ thống chưa tích hợp cổng thanh toán trực tuyến thật như VNPay, Momo, ZaloPay hoặc thẻ ngân hàng.",
        "Hệ thống chưa kết nối đơn vị vận chuyển thật nên chưa có mã vận đơn, phí vận chuyển động hoặc theo dõi giao hàng thời gian thực.",
        "Mật khẩu hiện được lưu ở mức cơ bản, chưa mã hóa bằng bcrypt và chưa có cơ chế JWT/session bảo vệ API phía backend.",
        "Phân quyền chủ yếu được kiểm tra ở frontend, khi triển khai thật cần kiểm tra quyền ở từng API backend.",
        "Chức năng ảnh sản phẩm sử dụng dữ liệu ảnh dạng base64, có thể làm dữ liệu lớn nếu số lượng sản phẩm nhiều.",
        "Chưa có kiểm thử tự động cho backend API và giao diện người dùng.",
        "Chưa có chức năng đánh giá, bình luận, xếp hạng sản phẩm hoặc yêu thích sản phẩm.",
        "Dashboard thống kê mới ở mức cơ bản, chưa có phân tích lợi nhuận, doanh thu theo thương hiệu hoặc tồn kho nâng cao.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3.6. Hướng phát triển", 3)
    add_para(
        doc,
        "Trong tương lai, website TAF Soccer có thể tiếp tục được hoàn thiện và mở rộng thêm nhiều chức năng nhằm nâng cao tính thực tế, tăng trải nghiệm người dùng và hỗ trợ tốt hơn cho công tác quản lý bán hàng.",
    )
    for item in [
        "Tích hợp cổng thanh toán trực tuyến thật để khách hàng có thể thanh toán ngay trên website.",
        "Kết nối với đơn vị vận chuyển để tạo mã vận đơn, tính phí vận chuyển và theo dõi trạng thái giao hàng.",
        "Mã hóa mật khẩu bằng bcrypt, bổ sung JWT/refresh token, xác minh email, quên mật khẩu và giới hạn đăng nhập sai.",
        "Bảo vệ các API quản trị bằng middleware kiểm tra token và role ở backend.",
        "Tách việc lưu ảnh sản phẩm sang upload file hoặc dịch vụ lưu trữ ảnh để giảm dung lượng document trong MongoDB.",
        "Bổ sung chức năng đánh giá sản phẩm, bình luận, xếp hạng sao và danh sách sản phẩm yêu thích.",
        "Bổ sung mã giảm giá, chương trình khuyến mãi và chính sách khách hàng thân thiết.",
        "Nâng cấp dashboard thống kê theo thương hiệu, theo size, theo khoảng thời gian và cảnh báo tồn kho thấp.",
        "Tối ưu giao diện responsive trên thiết bị di động, đặc biệt là giỏ hàng, form đặt hàng và trang quản trị.",
        "Triển khai website lên hosting hoặc cloud, cấu hình tên miền, HTTPS, sao lưu dữ liệu và môi trường production ổn định.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3.7. Kết luận", 3)
    add_para(
        doc,
        "Sau quá trình phân tích, thiết kế và cài đặt, đề tài “Xây dựng website bán giày bóng đá TAF Soccer” đã hoàn thành các chức năng chính của một website bán hàng trực tuyến. Hệ thống cho phép khách hàng xem danh sách sản phẩm, tìm kiếm, xem chi tiết, chọn size, thêm sản phẩm vào giỏ hàng, đặt hàng và theo dõi lịch sử đơn.",
    )
    add_para(
        doc,
        "Về mặt kỹ thuật, đề tài đã vận dụng HTML, CSS và JavaScript để xây dựng giao diện; sử dụng Node.js và ExpressJS để xây dựng backend; sử dụng MongoDB và Mongoose để lưu trữ, truy vấn và quản lý dữ liệu. Cách tổ chức này giúp hệ thống có cấu trúc rõ ràng, dễ hiểu và phù hợp với mô hình phát triển ứng dụng web hiện đại.",
    )
    add_para(
        doc,
        "Về mặt nghiệp vụ, website đã mô phỏng được quy trình bán giày bóng đá từ lúc khách hàng xem sản phẩm đến khi tạo đơn hàng. Hệ thống cũng hỗ trợ quản trị dữ liệu, phân quyền người dùng, quản lý sản phẩm, quản lý thương hiệu, quản lý đơn hàng và thống kê doanh số. Đây là các chức năng cần thiết đối với một website thương mại điện tử cơ bản.",
    )
    add_para(
        doc,
        "Thông qua đề tài, sinh viên rèn luyện được khả năng khảo sát yêu cầu, phân tích chức năng, thiết kế hệ thống, xây dựng cơ sở dữ liệu, lập trình API, tích hợp frontend - backend và kiểm thử sản phẩm. Mặc dù hệ thống vẫn còn một số hạn chế như chưa tích hợp thanh toán thật, chưa có vận chuyển thật và chưa hoàn thiện bảo mật nâng cao, nhưng đề tài đã đạt được mục tiêu học tập và là nền tảng để tiếp tục phát triển trong tương lai.",
    )

    doc.core_properties.title = "Chương 3 - Báo cáo TAF Soccer"
    doc.core_properties.author = "Ngô Khuê Văn"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_chapter3()
