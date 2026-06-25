from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Bao_cao_TAF_Soccer.docx"
ASSET_DIR = ROOT / "report_assets"

FONT = "Times New Roman"
FONT_DIAGRAM = "Arial"
INK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(80, 80, 80)


def set_run_font(run, name=FONT, size=13, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_base(p, after=6, before=0, line=1.15, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_para(doc, text="", bold=False, italic=False, size=13, align=None, after=6, before=0):
    p = doc.add_paragraph()
    set_paragraph_base(p, after=after, before=before, align=align)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size = 16
        color = INK
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size = 13
        color = INK
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size = 13
        color = INK
    else:
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(4)
        size = 12.5
        color = DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def add_caption(doc, text):
    p = add_para(doc, text, italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, before=2)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=12, align=None, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)


def set_table_geometry(table, widths_dxa, indent_dxa=120, cell_margin=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in {
        "top": 80,
        "bottom": 80,
        "start": cell_margin,
        "end": cell_margin,
    }.items():
        el = tbl_cell_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths_dxa=None, font_size=11.5, header_fill="F2F4F7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths_dxa is None:
        width = int(9360 / len(headers))
        widths_dxa = [width] * len(headers)
    set_table_geometry(table, widths_dxa)

    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER, fill=header_fill)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(row) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=font_size, align=align)
    add_para(doc, "", after=4)
    return table


def add_use_case(doc, name, actor, goal, pre, flow, result):
    add_para(doc, name, bold=True, after=2)
    add_table(
        doc,
        ["Mục", "Nội dung"],
        [
            ["Tên UC", name],
            ["Mục đích", goal],
            ["Tác nhân", actor],
            ["Tiền điều kiện", pre],
            ["Luồng xử lý chính", flow],
            ["Kết quả", result],
        ],
        widths_dxa=[2200, 7160],
        font_size=11.5,
    )


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=11)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(13)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, INK),
        ("Heading 2", 13, 12, 6, INK),
        ("Heading 3", 13, 8, 4, INK),
        ("Heading 4", 12.5, 7, 4, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_p.add_run("Báo cáo bài tập lớn - TAF Soccer")
    set_run_font(r, size=10.5, italic=True, color=GRAY)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


def load_font(name="arial.ttf", size=28):
    if ImageFont is None:
        return None
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/times.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_wrapped(draw, text, box, font, fill="#111111", align="center", line_gap=6):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 18
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_h = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    total_h = len(lines) * line_h
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textbbox((0, 0), line, font=font)[2]
        if align == "center":
            x = x1 + ((x2 - x1) - w) / 2
        else:
            x = x1 + 12
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def draw_box(draw, box, text, fill, outline="#344054", font=None, text_fill="#111111"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    draw_wrapped(draw, text, box, font, fill=text_fill)


def draw_arrow(draw, start, end, fill="#344054", width=4):
    import math

    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    spread = 0.55
    p1 = (
        end[0] - length * math.cos(angle - spread),
        end[1] - length * math.sin(angle - spread),
    )
    p2 = (
        end[0] - length * math.cos(angle + spread),
        end[1] - length * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=fill)


def save_diagram_architecture(path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font("arialbd.ttf", 38)
    font = load_font("arial.ttf", 25)
    small = load_font("arial.ttf", 21)
    draw.text((70, 45), "Kiến trúc tổng thể hệ thống TAF Soccer", font=title_font, fill="#0B2545")
    boxes = [
        ((80, 190, 430, 430), "Frontend\nHTML, CSS, JavaScript\nTrang khách hàng, nhân viên, admin", "#E8F3FF"),
        ((575, 190, 925, 430), "Backend\nNode.js + Express\nREST API xử lý nghiệp vụ", "#F2F4F7"),
        ((1070, 190, 1420, 430), "Database\nMongoDB\nUser, Product, Cart, Order, Brand, Address", "#ECFDF3"),
    ]
    for box, text, fill in boxes:
        draw_box(draw, box, text, fill, font=font)
    draw_arrow(draw, (430, 310), (575, 310))
    draw_arrow(draw, (925, 310), (1070, 310))
    notes = [
        ((90, 560, 420, 730), "Giao diện gọi API bằng fetch, lưu phiên đăng nhập trong localStorage."),
        ((585, 560, 915, 730), "Server mở port 3000, phục vụ thư mục Frontend và gom route theo từng nhóm API."),
        ((1080, 560, 1410, 730), "Mongoose model định nghĩa dữ liệu và thao tác CRUD với từng collection."),
    ]
    for box, text in notes:
        draw.rounded_rectangle(box, radius=14, fill="#FFFFFF", outline="#D0D5DD", width=2)
        draw_wrapped(draw, text, box, small, fill="#344054", align="left")
    img.save(path)


def save_diagram_use_case(path):
    img = Image.new("RGB", (1500, 950), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font("arialbd.ttf", 38)
    font = load_font("arial.ttf", 23)
    draw.text((70, 45), "Biểu đồ use case tổng quát", font=title_font, fill="#0B2545")
    actors = [
        ((70, 190, 280, 280), "Khách hàng"),
        ((70, 580, 280, 670), "Nhân viên"),
        ((1220, 390, 1430, 480), "Admin"),
    ]
    for box, text in actors:
        draw_box(draw, box, text, "#FFF7E6", font=font)
    use_cases = [
        ((430, 150, 700, 230), "Đăng ký / đăng nhập"),
        ((430, 255, 700, 335), "Xem, tìm kiếm sản phẩm"),
        ((430, 360, 700, 440), "Xem chi tiết sản phẩm"),
        ((430, 465, 700, 545), "Quản lý giỏ hàng"),
        ((430, 570, 700, 650), "Đặt hàng"),
        ((430, 675, 700, 755), "Xem lịch sử đơn"),
        ((810, 225, 1090, 305), "Quản lý sản phẩm"),
        ((810, 330, 1090, 410), "Quản lý đơn hàng"),
        ((810, 435, 1090, 515), "Thống kê doanh số"),
        ((810, 540, 1090, 620), "Quản lý người dùng"),
        ((810, 645, 1090, 725), "Quản lý thương hiệu"),
    ]
    for y in [190, 295, 400, 505, 610, 715]:
        draw.line([(280, 235), (430, y)], fill="#667085", width=3)
    for y in [265, 370, 475]:
        draw.line([(280, 625), (810, y)], fill="#667085", width=3)
    for y in [475, 580, 685]:
        draw.line([(1220, 435), (1090, y)], fill="#667085", width=3)
    for box, text in use_cases:
        draw_box(draw, box, text, "#EEF4FF", font=font)
    for box, text in actors:
        draw_box(draw, box, text, "#FFF7E6", font=font)
    img.save(path)


def save_diagram_sequence(path):
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font("arialbd.ttf", 38)
    font = load_font("arial.ttf", 22)
    small = load_font("arial.ttf", 20)
    draw.text((70, 45), "Biểu đồ tuần tự: khách hàng đặt hàng", font=title_font, fill="#0B2545")
    xs = [170, 500, 820, 1130, 1420]
    labels = ["Khách hàng", "Frontend cart.js", "API /api/orders", "MongoDB", "API /api/cart"]
    for x, label in zip(xs, labels):
        draw_box(draw, (x - 115, 145, x + 115, 225), label, "#F2F4F7", font=font)
        draw.line([(x, 225), (x, 840)], fill="#98A2B3", width=3)
    steps = [
        (170, 500, 300, "1. Chọn sản phẩm trong giỏ và nhập địa chỉ"),
        (500, 820, 390, "2. POST /api/orders kèm userId, products, totalPrice"),
        (820, 1130, 480, "3. Tạo bản ghi Order"),
        (820, 1420, 570, "4. Xóa sản phẩm đã đặt khỏi Cart"),
        (1420, 1130, 660, "5. Cập nhật giỏ hàng"),
        (820, 500, 750, "6. Trả kết quả đặt hàng thành công"),
    ]
    for x1, x2, y, text in steps:
        draw_arrow(draw, (x1, y), (x2, y))
        tx = min(x1, x2) + 20
        draw.text((tx, y - 32), text, font=small, fill="#344054")
    img.save(path)


def save_diagram_activity(path):
    img = Image.new("RGB", (1300, 1500), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font("arialbd.ttf", 38)
    font = load_font("arial.ttf", 24)
    draw.text((70, 45), "Biểu đồ hoạt động: quy trình mua hàng", font=title_font, fill="#0B2545")
    boxes = [
        ((390, 150, 910, 225), "Bắt đầu: khách hàng đăng nhập"),
        ((390, 290, 910, 365), "Xem danh sách và chi tiết sản phẩm"),
        ((390, 430, 910, 505), "Chọn size, số lượng và thêm vào giỏ"),
        ((390, 570, 910, 645), "Chọn sản phẩm cần thanh toán"),
        ((390, 710, 910, 785), "Nhập hoặc chọn địa chỉ nhận hàng"),
        ((390, 850, 910, 925), "Frontend gửi yêu cầu tạo đơn hàng"),
        ((390, 990, 910, 1065), "Backend lưu Order và cập nhật Cart"),
        ((390, 1130, 910, 1205), "Hiển thị thông báo đặt hàng thành công"),
        ((390, 1270, 910, 1345), "Kết thúc: đơn hàng chờ xác nhận"),
    ]
    for i, (box, text) in enumerate(boxes):
        draw_box(draw, box, text, "#ECFDF3" if i in [0, len(boxes) - 1] else "#EEF4FF", font=font)
        if i < len(boxes) - 1:
            draw_arrow(draw, ((box[0] + box[2]) // 2, box[3]), ((box[0] + box[2]) // 2, boxes[i + 1][0][1]))
    img.save(path)


def create_diagrams():
    if Image is None:
        return {}
    ASSET_DIR.mkdir(exist_ok=True)
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "use_case": ASSET_DIR / "use_case.png",
        "sequence_checkout": ASSET_DIR / "sequence_checkout.png",
        "activity_order": ASSET_DIR / "activity_order.png",
    }
    save_diagram_architecture(paths["architecture"])
    save_diagram_use_case(paths["use_case"])
    save_diagram_sequence(paths["sequence_checkout"])
    save_diagram_activity(paths["activity_order"])
    return paths


def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360], indent_dxa=0)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = [
        ("TRƯỜNG ĐẠI HỌC", 14, True),
        ("TÀI CHÍNH - NGÂN HÀNG HÀ NỘI", 14, True),
        ("VIỆN CÔNG NGHỆ THÔNG TIN", 14, True),
        ("", 12, False),
        ("BÀI TẬP LỚN", 19, True),
        ("BỘ MÔN JAVASCRIPT VÀ LẬP TRÌNH WEB", 14, True),
        ("", 12, False),
        ("ĐỀ TÀI", 14, True),
        ("XÂY DỰNG WEBSITE BÁN GIÀY BÓNG ĐÁ", 18, True),
        ("TAF SOCCER", 22, True),
        ("", 12, False),
        ("Họ tên SV:        NGÔ KHUÊ VĂN", 13, False),
        ("MSSV:             2354800053", 13, False),
        ("Khóa:             D12.48.02", 13, False),
        ("Giảng viên hướng dẫn: VŨ THỊ THANH HUYỀN", 13, False),
        ("", 12, False),
        ("Hà Nội, ngày 13 tháng 6 năm 2026", 13, False),
    ]
    first = True
    for text, size, bold in lines:
        if first:
            first = False
        else:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5 if text else 12)
        if text:
            r = p.add_run(text)
            set_run_font(r, size=size, bold=bold)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "MỤC LỤC", 1)
    items = [
        "LỜI MỞ ĐẦU",
        "CHƯƠNG 1. VẤN ĐỀ VÀ CƠ SỞ LÝ THUYẾT",
        "1.1 Lý do lựa chọn đề tài",
        "1.2 Tầm quan trọng của đề tài",
        "1.3 Mục tiêu và nội dung thực hiện",
        "1.4 Đối tượng, phạm vi và phương pháp",
        "1.5 Công cụ và công nghệ sử dụng",
        "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "2.1 Phân tích yêu cầu chức năng",
        "2.2 Cấu trúc logic, use case và luồng xử lý",
        "2.3 Thiết kế cơ sở dữ liệu và API",
        "CHƯƠNG 3. CÀI ĐẶT, ĐÁNH GIÁ VÀ KẾT LUẬN",
        "3.1 Cài đặt hệ thống",
        "3.2 Lập trình và tích hợp",
        "3.3 Kiểm thử, đánh giá và kết luận",
        "TÀI LIỆU THAM KHẢO",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.45 if item[0].isdigit() else 0)
        r = p.add_run(item)
        set_run_font(r, size=13, bold=item.startswith("CHƯƠNG") or item in ["LỜI MỞ ĐẦU", "TÀI LIỆU THAM KHẢO"])
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "LỜI MỞ ĐẦU", 1)
    paragraphs = [
        "Trong thời đại công nghệ thông tin phát triển mạnh mẽ, website thương mại điện tử đã trở thành công cụ quan trọng giúp cửa hàng giới thiệu sản phẩm, tiếp cận khách hàng và quản lý hoạt động bán hàng hiệu quả hơn. Đối với lĩnh vực thể thao, đặc biệt là giày bóng đá, nhu cầu tìm kiếm sản phẩm theo thương hiệu, size, giá bán và tình trạng tồn kho ngày càng rõ rệt.",
        "Trong quá trình học tập môn JavaScript và Lập trình Web, em đã được tiếp cận với các kiến thức về HTML, CSS, JavaScript, Node.js, ExpressJS và MongoDB. Trên cơ sở đó, em thực hiện đề tài xây dựng website bán giày bóng đá TAF Soccer nhằm vận dụng kiến thức đã học vào một hệ thống có chức năng thực tế.",
        "Bài tập lớn hướng tới việc xây dựng một website có đầy đủ các chức năng cơ bản: đăng ký, đăng nhập, hiển thị sản phẩm, xem chi tiết sản phẩm, quản lý giỏ hàng, đặt hàng, quản lý sản phẩm, quản lý thương hiệu, quản lý đơn hàng, quản lý người dùng và thống kê doanh số.",
        "Do kiến thức và kinh nghiệm thực tế còn hạn chế, bài báo cáo khó tránh khỏi thiếu sót. Em rất mong nhận được ý kiến đóng góp của cô để có thể hoàn thiện sản phẩm và nâng cao kỹ năng lập trình web trong các dự án tiếp theo.",
        "Em xin chân thành cảm ơn cô Vũ Thị Thanh Huyền đã tận tình giảng dạy và hướng dẫn em trong quá trình học tập cũng như thực hiện bài tập lớn này.",
    ]
    for text in paragraphs:
        add_para(doc, text)
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. VẤN ĐỀ VÀ CƠ SỞ LÝ THUYẾT", 1)
    add_heading(doc, "1.1 Lý do lựa chọn đề tài", 2)
    for text in [
        "Giày bóng đá là mặt hàng có nhu cầu cao đối với học sinh, sinh viên và người chơi thể thao phong trào. Khách hàng thường cần tham khảo nhiều thông tin trước khi mua như thương hiệu, mẫu mã, size, giá, hình ảnh và tình trạng còn hàng.",
        "Nhiều cửa hàng vẫn quản lý sản phẩm và đơn hàng thủ công, gây khó khăn khi số lượng mẫu giày, size và đơn hàng tăng lên. Việc xây dựng website bán giày bóng đá giúp cửa hàng số hóa quy trình bán hàng, đồng thời giúp khách hàng dễ dàng tra cứu và đặt mua sản phẩm trực tuyến.",
        "Đề tài TAF Soccer được lựa chọn vì vừa gần gũi với nhu cầu thực tế, vừa phù hợp để vận dụng kiến thức lập trình web: xây dựng giao diện, xử lý tương tác bằng JavaScript, tạo REST API bằng ExpressJS và lưu trữ dữ liệu bằng MongoDB.",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.2 Tầm quan trọng của đề tài", 2)
    add_para(doc, "Việc xây dựng website bán giày bóng đá TAF Soccer có ý nghĩa cả về mặt thực tiễn và học tập.")
    add_para(doc, "Về mặt thực tiễn, website hỗ trợ cửa hàng:")
    for item in [
        "Quảng bá sản phẩm đến nhiều khách hàng hơn thông qua môi trường trực tuyến.",
        "Quản lý danh sách sản phẩm, thương hiệu, size và số lượng tồn kho.",
        "Tiếp nhận đơn hàng, cập nhật trạng thái giao hàng và theo dõi thanh toán.",
        "Tổng hợp doanh số, sản phẩm bán chạy và trạng thái đơn hàng để phục vụ quản lý.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Đối với khách hàng, hệ thống mang lại các tiện ích:")
    for item in [
        "Tìm kiếm và xem thông tin chi tiết sản phẩm mọi lúc.",
        "Chọn size phù hợp trước khi thêm sản phẩm vào giỏ hàng.",
        "Đặt hàng trực tuyến và lưu thông tin địa chỉ nhận hàng cho các lần mua tiếp theo.",
        "Theo dõi lịch sử đơn hàng và trạng thái xử lý đơn.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Mục tiêu và nội dung thực hiện", 2)
    add_heading(doc, "Mục tiêu của đề tài", 3)
    add_para(doc, "Đề tài hướng tới xây dựng một website thương mại điện tử bán giày bóng đá với các chức năng cơ bản, giao diện dễ sử dụng và có khả năng quản lý dữ liệu bằng cơ sở dữ liệu MongoDB.")
    for item in [
        "Thiết kế giao diện người dùng cho khách hàng, nhân viên và admin.",
        "Xây dựng chức năng đăng ký, đăng nhập và phân quyền theo vai trò.",
        "Hiển thị danh sách sản phẩm, lọc/tìm kiếm sản phẩm và xem chi tiết.",
        "Quản lý giỏ hàng theo từng tài khoản, size và số lượng.",
        "Tạo đơn hàng, lưu địa chỉ nhận hàng và quản lý trạng thái đơn.",
        "Quản lý sản phẩm, thương hiệu, người dùng và thống kê doanh số.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Nội dung thực hiện", 3)
    for item in [
        "Khảo sát nhu cầu của website bán giày bóng đá và xác định chức năng chính.",
        "Thiết kế giao diện bằng HTML, CSS và JavaScript thuần.",
        "Xây dựng backend bằng Node.js, ExpressJS và các route RESTful API.",
        "Thiết kế các model dữ liệu bằng Mongoose và lưu trữ trong MongoDB.",
        "Tích hợp frontend với backend bằng fetch API.",
        "Kiểm thử chức năng, ghi nhận kết quả và đề xuất hướng phát triển.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.4 Đối tượng, phạm vi và phương pháp", 2)
    add_heading(doc, "Đối tượng nghiên cứu", 3)
    for item in [
        "Website thương mại điện tử bán giày bóng đá TAF Soccer.",
        "Người dùng có nhu cầu xem, tìm kiếm và mua giày bóng đá trực tuyến.",
        "Nhân viên bán hàng cần quản lý sản phẩm, đơn hàng và thống kê.",
        "Admin cần quản lý tài khoản người dùng và thương hiệu sản phẩm.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Phạm vi nghiên cứu", 3)
    add_para(doc, "Đề tài tập trung vào các chức năng cốt lõi của một website bán hàng:")
    for item in [
        "Đăng ký, đăng nhập và lưu phiên người dùng ở trình duyệt.",
        "Hiển thị sản phẩm theo danh sách và theo thương hiệu.",
        "Quản lý giỏ hàng, chọn sản phẩm thanh toán và đặt hàng.",
        "Quản lý sản phẩm có tồn kho theo size.",
        "Quản lý thương hiệu, người dùng, đơn hàng và thống kê doanh số.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Đề tài chưa tập trung vào các chức năng nâng cao như cổng thanh toán online, gửi email tự động, mã hóa mật khẩu bằng bcrypt/JWT, tích hợp vận chuyển thực tế hoặc hệ thống gợi ý sản phẩm.")
    add_heading(doc, "Phương pháp thực hiện", 3)
    for item in [
        "Phương pháp khảo sát và phân tích yêu cầu chức năng.",
        "Phương pháp thiết kế hệ thống theo mô hình client-server.",
        "Phương pháp lập trình từng module, sau đó tích hợp và kiểm thử.",
        "Phương pháp đọc mã nguồn để mô tả lại API, model dữ liệu và luồng xử lý.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.5 Công cụ và công nghệ sử dụng", 2)
    tech_rows = [
        ["HTML", "Xây dựng cấu trúc các trang giao diện như trang chủ, đăng nhập, sản phẩm, giỏ hàng, quản trị."],
        ["CSS", "Định dạng bố cục, màu sắc, bảng, form, dashboard và khả năng hiển thị thân thiện."],
        ["JavaScript", "Xử lý sự kiện, gọi API bằng fetch, kiểm tra dữ liệu form và cập nhật giao diện động."],
        ["Node.js", "Môi trường chạy backend JavaScript."],
        ["ExpressJS", "Xây dựng REST API, route xử lý xác thực, sản phẩm, giỏ hàng, đơn hàng, thống kê."],
        ["MongoDB", "Lưu trữ dữ liệu người dùng, sản phẩm, giỏ hàng, đơn hàng, thương hiệu và địa chỉ."],
        ["Mongoose", "Định nghĩa schema/model và thao tác dữ liệu MongoDB trong backend."],
        ["Visual Studio Code", "Môi trường viết và quản lý mã nguồn dự án."],
    ]
    add_table(doc, ["Công nghệ", "Vai trò trong hệ thống"], tech_rows, widths_dxa=[2300, 7060])
    doc.add_page_break()


def add_chapter_2(doc, diagrams):
    add_heading(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "2.1 Phân tích yêu cầu, chức năng hệ thống", 2)
    add_heading(doc, "2.1.1 Yêu cầu hệ thống", 3)
    add_para(doc, "Website TAF Soccer được xây dựng nhằm hỗ trợ khách hàng tìm kiếm, lựa chọn và đặt mua giày bóng đá trực tuyến; đồng thời hỗ trợ nhân viên và admin quản lý dữ liệu bán hàng.")
    role_rows = [
        ["Khách hàng", "Đăng ký, đăng nhập, xem sản phẩm, tìm kiếm, xem chi tiết, thêm vào giỏ, đặt hàng, lưu địa chỉ, xem lịch sử đơn hàng, đổi mật khẩu."],
        ["Nhân viên", "Đăng nhập trang quản trị nhân viên, quản lý sản phẩm, quản lý đơn hàng, cập nhật trạng thái giao hàng, đánh dấu thanh toán, xem dashboard và thống kê doanh số."],
        ["Admin", "Quản lý tài khoản người dùng, phân quyền, quản lý thương hiệu và kiểm soát dữ liệu hệ thống."],
    ]
    add_table(doc, ["Tác nhân", "Yêu cầu chức năng"], role_rows, widths_dxa=[1900, 7460])

    add_heading(doc, "2.1.2 Các yêu cầu phi chức năng", 3)
    nfr_rows = [
        ["Hiệu năng", "Các thao tác xem sản phẩm, tìm kiếm, thêm giỏ hàng và cập nhật đơn hàng cần phản hồi nhanh trong môi trường thử nghiệm cục bộ."],
        ["Tính khả dụng", "Giao diện chia rõ khu vực khách hàng, nhân viên và admin; thao tác form, bảng, nút thêm/sửa/xóa dễ nhận biết."],
        ["Bảo mật", "Hệ thống có phân quyền theo role trong frontend. Mật khẩu hiện đang lưu dạng chuỗi thường, nên cần nâng cấp mã hóa và token xác thực khi triển khai thực tế."],
        ["Bảo trì", "Mã nguồn tách thành Frontend, Backend, routes, models và CSS/JS theo từng trang, giúp dễ mở rộng chức năng."],
        ["Khả năng mở rộng", "Có thể bổ sung thanh toán online, email xác nhận, quản lý khuyến mãi, đánh giá sản phẩm và API vận chuyển."],
    ]
    add_table(doc, ["Yêu cầu", "Mô tả"], nfr_rows, widths_dxa=[1900, 7460])

    add_heading(doc, "2.1.3 Các chức năng chính của hệ thống", 3)
    function_groups = [
        ("Chức năng người dùng", ["Đăng ký tài khoản", "Đăng nhập, đăng xuất", "Đổi mật khẩu", "Xem lịch sử đơn hàng"]),
        ("Chức năng sản phẩm", ["Xem danh sách sản phẩm", "Tìm kiếm theo tên hoặc thương hiệu", "Xem chi tiết sản phẩm", "Chọn size và số lượng"]),
        ("Chức năng giỏ hàng", ["Thêm sản phẩm vào giỏ", "Cập nhật số lượng", "Xóa một/nhiều sản phẩm", "Tính tổng tiền và phí vận chuyển"]),
        ("Chức năng đặt hàng", ["Nhập thông tin nhận hàng", "Lưu địa chỉ nhận hàng", "Tạo đơn hàng", "Xóa sản phẩm đã đặt khỏi giỏ"]),
        ("Chức năng quản trị", ["Quản lý sản phẩm", "Quản lý thương hiệu", "Quản lý người dùng", "Quản lý đơn hàng", "Thống kê doanh số"]),
    ]
    for title, items in function_groups:
        add_para(doc, title, bold=True, after=2)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "2.2 Phân tích cấu trúc logic hệ thống", 2)
    add_heading(doc, "2.2.1 Biểu đồ use case", 3)
    if diagrams.get("use_case"):
        doc.add_picture(str(diagrams["use_case"]), width=Inches(6.4))
        add_caption(doc, "Hình 2.1. Biểu đồ use case tổng quát của hệ thống")

    add_heading(doc, "2.2.1.1 Phân rã use case", 4)
    use_cases = [
        ("Đăng ký/Đăng nhập", "Người dùng", "Cho phép người dùng tạo tài khoản và truy cập hệ thống theo vai trò.", "Người dùng có email và mật khẩu hợp lệ.", "Người dùng nhập thông tin; frontend gửi request đến /api/auth; backend kiểm tra dữ liệu và trả thông tin tài khoản.", "Người dùng đăng nhập thành công và được chuyển đến giao diện phù hợp."),
        ("Xem và tìm kiếm sản phẩm", "Khách hàng", "Cho phép khách hàng xem danh sách giày bóng đá và tìm kiếm sản phẩm phù hợp.", "Backend đang chạy và có dữ liệu sản phẩm.", "Frontend gọi /api/products hoặc /api/products/search/:keyword; dữ liệu được hiển thị thành danh sách sản phẩm.", "Danh sách sản phẩm phù hợp được hiển thị."),
        ("Xem chi tiết sản phẩm", "Khách hàng", "Hiển thị thông tin chi tiết, thương hiệu, giá, hình ảnh, mô tả, size và số lượng tồn.", "Sản phẩm tồn tại trong cơ sở dữ liệu.", "Người dùng chọn sản phẩm; frontend lưu selectedProductId và gọi /api/products/:id.", "Trang chi tiết hiển thị đúng dữ liệu sản phẩm."),
        ("Quản lý giỏ hàng", "Khách hàng", "Cho phép thêm, cập nhật, xóa và chọn sản phẩm cần thanh toán.", "Người dùng đã đăng nhập bằng tài khoản khách hàng.", "Frontend gọi các API /api/cart/add, /api/cart/update, /api/cart/remove và /api/cart/clear/:userId.", "Giỏ hàng được cập nhật đúng theo thao tác của người dùng."),
        ("Đặt hàng", "Khách hàng", "Tạo đơn hàng từ các sản phẩm được chọn trong giỏ hàng.", "Giỏ hàng có ít nhất một sản phẩm và người dùng nhập đầy đủ địa chỉ nhận hàng.", "Frontend gửi POST /api/orders; backend tạo Order và xóa sản phẩm đã đặt khỏi Cart.", "Đơn hàng được lưu và chuyển sang trạng thái Chờ xác nhận."),
        ("Quản lý sản phẩm", "Nhân viên", "Cho phép nhân viên thêm, sửa, xóa và tìm kiếm sản phẩm.", "Nhân viên đăng nhập và truy cập trang staff-products.", "Nhân viên nhập thông tin sản phẩm, thương hiệu, giá, ảnh, mô tả và tồn kho theo size; frontend gọi /api/products.", "Dữ liệu sản phẩm được cập nhật trong MongoDB."),
        ("Quản lý đơn hàng", "Nhân viên", "Theo dõi đơn hàng, cập nhật trạng thái và đánh dấu thanh toán.", "Có đơn hàng trong hệ thống.", "Nhân viên xem danh sách đơn; cập nhật trạng thái theo luồng Chờ xác nhận -> Đã xác nhận -> Đang giao -> Đã giao -> Hoàn thành.", "Trạng thái đơn hàng được cập nhật theo đúng quy tắc nghiệp vụ."),
        ("Quản lý người dùng và thương hiệu", "Admin", "Admin quản lý tài khoản, phân quyền và danh sách thương hiệu.", "Admin đăng nhập hệ thống.", "Frontend admin gọi /api/admin/users và /api/admin/brands để thêm/sửa/xóa/tìm kiếm.", "Tài khoản và thương hiệu được quản lý tập trung."),
        ("Thống kê doanh số", "Nhân viên/Admin", "Hiển thị tổng sản phẩm, tổng đơn, doanh thu, sản phẩm bán chạy và trạng thái đơn hàng.", "Có dữ liệu đơn hàng trong hệ thống.", "Frontend gọi /api/statistics/sales, có thể truyền from/to để lọc theo ngày.", "Dashboard thống kê hiển thị số liệu phục vụ quản lý."),
    ]
    for uc in use_cases:
        add_use_case(doc, *uc)

    add_heading(doc, "2.2.2 Biểu đồ tuần tự", 3)
    if diagrams.get("sequence_checkout"):
        doc.add_picture(str(diagrams["sequence_checkout"]), width=Inches(6.4))
        add_caption(doc, "Hình 2.2. Biểu đồ tuần tự chức năng đặt hàng")
    add_para(doc, "Luồng đặt hàng bắt đầu từ trang giỏ hàng. Khách hàng chọn sản phẩm cần thanh toán, nhập thông tin nhận hàng, sau đó frontend gửi dữ liệu sang API tạo đơn hàng. Backend kiểm tra thông tin, lưu đơn hàng vào MongoDB và xóa các sản phẩm đã đặt khỏi giỏ hàng của người dùng.")

    add_heading(doc, "2.2.3 Biểu đồ hoạt động", 3)
    if diagrams.get("activity_order"):
        doc.add_picture(str(diagrams["activity_order"]), width=Inches(5.5))
        add_caption(doc, "Hình 2.3. Biểu đồ hoạt động quy trình mua hàng")

    add_heading(doc, "2.2.4 Cấu trúc logic hệ thống", 3)
    if diagrams.get("architecture"):
        doc.add_picture(str(diagrams["architecture"]), width=Inches(6.4))
        add_caption(doc, "Hình 2.4. Kiến trúc tổng thể hệ thống TAF Soccer")
    add_para(doc, "Hệ thống được thiết kế theo mô hình client-server và có thể mô tả theo ba lớp chính:")
    layers = [
        ("Lớp giao diện", "Được xây dựng bằng HTML, CSS, JavaScript. Lớp này hiển thị dữ liệu sản phẩm, form đăng nhập/đăng ký, giỏ hàng, đơn hàng, dashboard nhân viên và trang admin."),
        ("Lớp xử lý nghiệp vụ", "Backend Node.js/Express tiếp nhận request, kiểm tra dữ liệu đầu vào, xử lý nghiệp vụ sản phẩm, giỏ hàng, đơn hàng, thương hiệu, người dùng và thống kê."),
        ("Lớp dữ liệu", "MongoDB lưu trữ các collection User, Product, Cart, Order, Brand và Address; Mongoose được dùng để định nghĩa schema và thao tác dữ liệu."),
    ]
    add_table(doc, ["Lớp", "Mô tả"], layers, widths_dxa=[2000, 7360])

    add_heading(doc, "2.3 Thiết kế cơ sở dữ liệu và API", 2)
    add_heading(doc, "2.3.1 Thiết kế cơ sở dữ liệu", 3)
    db_tables = [
        ("Bảng User", [["_id", "ObjectId", "Mã người dùng"], ["name", "String", "Họ tên"], ["email", "String", "Email đăng nhập, duy nhất"], ["password", "String", "Mật khẩu"], ["phone", "String", "Số điện thoại"], ["role", "String", "Vai trò: user, staff, admin"], ["createdAt/updatedAt", "Date", "Thời gian tạo/cập nhật"]]),
        ("Bảng Product", [["_id", "ObjectId", "Mã sản phẩm"], ["name", "String", "Tên sản phẩm"], ["brand", "String", "Thương hiệu"], ["price", "Number", "Giá bán"], ["sizeStock", "Array", "Danh sách size và số lượng tồn"], ["quantity", "Number", "Tổng số lượng tồn"], ["image", "String", "Ảnh sản phẩm"], ["description", "String", "Mô tả sản phẩm"]]),
        ("Bảng Cart", [["_id", "ObjectId", "Mã giỏ hàng"], ["userId", "String", "Mã người dùng"], ["products", "Array", "Danh sách sản phẩm trong giỏ"], ["productId", "String", "Mã sản phẩm"], ["size", "Number", "Size đã chọn"], ["quantity", "Number", "Số lượng trong giỏ"]]),
        ("Bảng Order", [["_id", "ObjectId", "Mã đơn hàng"], ["userId", "String", "Mã người đặt"], ["customerName", "String", "Tên người nhận"], ["phone", "String", "Số điện thoại"], ["address", "String", "Địa chỉ giao hàng"], ["products", "Array", "Danh sách sản phẩm đặt"], ["totalPrice", "Number", "Tổng tiền"], ["status", "String", "Trạng thái đơn hàng"], ["paymentStatus", "String", "Trạng thái thanh toán"], ["paidAt/createdAt", "Date", "Thời điểm thanh toán/tạo đơn"]]),
        ("Bảng Brand", [["_id", "ObjectId", "Mã thương hiệu"], ["name", "String", "Tên thương hiệu, duy nhất"], ["description", "String", "Mô tả thương hiệu"], ["createdAt/updatedAt", "Date", "Thời gian tạo/cập nhật"]]),
        ("Bảng Address", [["_id", "ObjectId", "Mã địa chỉ"], ["userId", "String", "Mã người dùng"], ["name", "String", "Tên người nhận"], ["phone", "String", "Số điện thoại"], ["province", "String", "Tỉnh/thành phố"], ["ward", "String", "Xã/phường"], ["detailAddress", "String", "Địa chỉ chi tiết"], ["address", "String", "Địa chỉ đầy đủ"], ["createdAt", "Date", "Thời gian lưu"]]),
    ]
    for title, rows in db_tables:
        add_para(doc, title, bold=True, after=2)
        add_table(doc, ["Tên trường", "Kiểu dữ liệu", "Mô tả"], rows, widths_dxa=[2300, 1900, 5160], font_size=11.3)

    add_heading(doc, "2.3.2 Thiết kế API", 3)
    api_rows = [
        ["Auth", "POST /api/auth/register", "Đăng ký tài khoản khách hàng."],
        ["Auth", "POST /api/auth/login", "Đăng nhập và trả về thông tin người dùng."],
        ["Auth", "PUT /api/auth/change-password", "Đổi mật khẩu tài khoản."],
        ["Products", "GET /api/products", "Lấy danh sách sản phẩm."],
        ["Products", "POST /api/products", "Thêm sản phẩm mới."],
        ["Products", "GET /api/products/:id", "Lấy chi tiết sản phẩm."],
        ["Products", "PUT /api/products/:id", "Cập nhật sản phẩm."],
        ["Products", "DELETE /api/products/:id", "Xóa sản phẩm."],
        ["Products", "GET /api/products/search/:keyword", "Tìm kiếm theo tên hoặc thương hiệu."],
        ["Cart", "GET /api/cart/:userId", "Lấy giỏ hàng của người dùng."],
        ["Cart", "POST /api/cart/add", "Thêm sản phẩm vào giỏ."],
        ["Cart", "PUT /api/cart/update", "Cập nhật số lượng sản phẩm trong giỏ."],
        ["Cart", "DELETE /api/cart/remove", "Xóa một sản phẩm khỏi giỏ."],
        ["Cart", "DELETE /api/cart/clear/:userId", "Làm trống giỏ hàng."],
        ["Orders", "POST /api/orders", "Tạo đơn hàng từ giỏ hàng."],
        ["Orders", "GET /api/orders/user/:userId", "Lấy lịch sử đơn của khách hàng."],
        ["Orders", "GET /api/orders", "Nhân viên lấy danh sách đơn hàng."],
        ["Orders", "PUT /api/orders/:id/status", "Cập nhật trạng thái đơn hàng."],
        ["Orders", "PUT /api/orders/:id/payment", "Đánh dấu đơn hàng đã thanh toán."],
        ["Admin", "GET/POST/PUT/DELETE /api/admin/users", "Quản lý tài khoản và phân quyền."],
        ["Admin", "GET/POST/PUT/DELETE /api/admin/brands", "Quản lý thương hiệu sản phẩm."],
        ["Address", "GET/POST/DELETE /api/addresses", "Lưu và quản lý địa chỉ nhận hàng."],
        ["Statistics", "GET /api/statistics/sales", "Thống kê doanh số, sản phẩm bán chạy và trạng thái đơn."],
    ]
    add_table(doc, ["Nhóm", "Đường dẫn", "Mục đích"], api_rows, widths_dxa=[1350, 3300, 4710], font_size=10.5)
    doc.add_page_break()


def add_chapter_3(doc):
    add_heading(doc, "CHƯƠNG 3. CÀI ĐẶT, ĐÁNH GIÁ VÀ KẾT LUẬN", 1)
    add_heading(doc, "3.1 Cài đặt, yêu cầu trang thiết bị và phần mềm hệ thống", 2)
    add_heading(doc, "3.1.1 Yêu cầu phần cứng", 3)
    for item in ["Máy tính hoặc laptop có kết nối Internet.", "Bộ nhớ RAM từ 4GB trở lên.", "Dung lượng ổ cứng còn trống tối thiểu 5GB.", "Bộ xử lý đủ khả năng chạy Node.js, MongoDB và trình duyệt kiểm thử."]:
        add_bullet(doc, item)
    add_heading(doc, "3.1.2 Yêu cầu phần mềm", 3)
    for item in ["Hệ điều hành Windows 10 hoặc Windows 11.", "Visual Studio Code để lập trình và quản lý mã nguồn.", "Node.js và npm để chạy backend.", "MongoDB để lưu trữ dữ liệu.", "Trình duyệt Chrome/Edge để kiểm thử giao diện.", "Công cụ StarUML/draw.io có thể dùng để vẽ UML nếu cần bổ sung hình ảnh." ]:
        add_bullet(doc, item)
    add_heading(doc, "3.1.3 Cấu trúc thư mục mã nguồn", 3)
    structure_rows = [
        ["Backend/server.js", "Khởi tạo Express, kết nối MongoDB taf_soccer, khai báo route API và phục vụ thư mục Frontend."],
        ["Backend/models", "Chứa các model User, Product, Cart, Order, Brand, Address."],
        ["Backend/routes", "Chứa route xử lý auth, product, cart, order, statistic, admin và address."],
        ["Frontend/user", "Các trang khách hàng: homepage, all-products, product-detail, cart, account, login, register."],
        ["Frontend/staff", "Các trang nhân viên: dashboard, products, orders, statistics."],
        ["Frontend/admin", "Các trang admin: quản lý users và brands."],
        ["Frontend/JS", "Mã JavaScript xử lý gọi API và tương tác từng trang."],
        ["Frontend/CSS", "Mã CSS định dạng giao diện."],
    ]
    add_table(doc, ["Thành phần", "Vai trò"], structure_rows, widths_dxa=[2600, 6760])

    add_heading(doc, "3.2 Lập trình và tích hợp hệ thống", 2)
    add_heading(doc, "3.2.1 Xây dựng giao diện hệ thống", 3)
    add_para(doc, "Frontend của TAF Soccer được xây dựng bằng HTML, CSS và JavaScript thuần. Các trang được chia theo vai trò người dùng nhằm giúp giao diện rõ ràng và dễ quản lý.")
    ui_rows = [
        ["Trang khách hàng", "Đăng nhập, đăng ký, trang chủ, tất cả sản phẩm, sản phẩm theo thương hiệu, chi tiết sản phẩm, giỏ hàng, tài khoản và lịch sử đơn hàng."],
        ["Trang nhân viên", "Dashboard đơn hàng mới, quản lý sản phẩm, quản lý đơn hàng, thống kê doanh số."],
        ["Trang admin", "Quản lý tài khoản người dùng, phân quyền và quản lý thương hiệu."],
    ]
    add_table(doc, ["Nhóm giao diện", "Các trang/chức năng"], ui_rows, widths_dxa=[2100, 7260])

    add_heading(doc, "3.2.2 Xây dựng backend", 3)
    add_para(doc, "Backend sử dụng ExpressJS, cho phép frontend gọi API qua địa chỉ http://localhost:3000. File server.js thiết lập cors, parser JSON giới hạn 10MB để hỗ trợ ảnh base64, phục vụ thư mục Frontend tĩnh và kết nối MongoDB với database taf_soccer.")
    backend_rows = [
        ["authRoutes", "Xử lý đăng ký, đăng nhập và đổi mật khẩu."],
        ["productRoutes", "CRUD sản phẩm, chuẩn hóa tồn kho theo size và tìm kiếm sản phẩm."],
        ["cartRoutes", "Tạo/lấy giỏ hàng, thêm/cập nhật/xóa sản phẩm, làm trống giỏ hàng."],
        ["orderRoutes", "Tạo đơn, lấy danh sách đơn, cập nhật trạng thái và thanh toán."],
        ["adminRoutes", "Quản lý người dùng, phân quyền, thương hiệu và đồng bộ thương hiệu với sản phẩm."],
        ["statisticRoutes", "Tính doanh thu, số đơn hoàn thành, top sản phẩm bán chạy và thống kê trạng thái đơn hàng."],
        ["addressRoutes", "Lưu, lấy và xóa địa chỉ nhận hàng của người dùng."],
    ]
    add_table(doc, ["Route", "Nhiệm vụ"], backend_rows, widths_dxa=[2100, 7260])

    add_heading(doc, "3.2.3 Xây dựng cơ sở dữ liệu", 3)
    add_para(doc, "Dữ liệu được lưu trong MongoDB và thao tác thông qua Mongoose model. Product có trường sizeStock để quản lý tồn kho theo từng size giày; Order lưu danh sách sản phẩm đã đặt, tổng tiền, trạng thái đơn hàng và trạng thái thanh toán; Address lưu các địa chỉ nhận hàng đã dùng.")
    add_para(doc, "Các collection chính gồm: User, Product, Cart, Order, Brand và Address.")

    add_heading(doc, "3.2.4 Tích hợp hệ thống", 3)
    add_para(doc, "Hệ thống hoạt động theo mô hình client-server. Frontend sử dụng fetch API để gửi request đến backend. Backend xử lý nghiệp vụ, đọc/ghi MongoDB và trả dữ liệu JSON cho frontend hiển thị.")
    add_heading(doc, "Luồng xử lý chính của hệ thống", 3)
    main_flow = [
        "Người dùng truy cập website và đăng nhập/đăng ký tài khoản.",
        "Frontend lưu thông tin người dùng trong localStorage để nhận biết vai trò.",
        "Khi xem sản phẩm, frontend gọi API /api/products hoặc /api/products/:id.",
        "Khi thêm vào giỏ, frontend gửi userId, productId, size và quantity đến /api/cart/add.",
        "Khi đặt hàng, frontend gửi danh sách sản phẩm đã chọn và địa chỉ nhận hàng đến /api/orders.",
        "Backend tạo đơn hàng mới, lưu vào MongoDB và xóa sản phẩm đã đặt khỏi giỏ hàng.",
        "Nhân viên theo dõi đơn hàng, cập nhật trạng thái và đánh dấu thanh toán.",
        "Trang thống kê tổng hợp doanh thu theo tháng, top sản phẩm và số đơn theo trạng thái.",
    ]
    for item in main_flow:
        add_number(doc, item)

    add_heading(doc, "3.3 Thử nghiệm, đánh giá hệ thống và kết luận", 2)
    add_heading(doc, "3.3.1 Kết quả thực hiện", 3)
    add_para(doc, "Sau quá trình xây dựng, hệ thống đã hoàn thiện các chức năng cơ bản của website bán giày bóng đá:")
    for item in [
        "Khách hàng có thể đăng ký, đăng nhập, xem sản phẩm, tìm kiếm, xem chi tiết, thêm giỏ hàng và đặt hàng.",
        "Giỏ hàng hỗ trợ chọn sản phẩm cần thanh toán, cập nhật số lượng, xóa từng sản phẩm hoặc xóa nhiều sản phẩm đã chọn.",
        "Đơn hàng lưu thông tin người nhận, địa chỉ, danh sách sản phẩm, tổng tiền, trạng thái đơn và trạng thái thanh toán.",
        "Nhân viên có thể quản lý sản phẩm, tồn kho theo size, ảnh sản phẩm, mô tả và đơn hàng.",
        "Admin có thể quản lý tài khoản, phân quyền và thương hiệu sản phẩm.",
        "Trang thống kê hiển thị doanh thu, số đơn hoàn thành, sản phẩm bán chạy và tình trạng đơn hàng.",
    ]:
        add_bullet(doc, item)

    add_para(doc, "Bảng kiểm thử một số chức năng chính của hệ thống:", bold=True)
    test_rows = [
        ["1", "Đăng ký tài khoản", "Nhập họ tên, email, mật khẩu hợp lệ", "Tài khoản được tạo", "Đạt"],
        ["2", "Đăng nhập", "Nhập email và mật khẩu đúng", "Chuyển vào giao diện theo vai trò", "Đạt"],
        ["3", "Xem danh sách sản phẩm", "Mở trang tất cả sản phẩm", "Danh sách sản phẩm hiển thị", "Đạt"],
        ["4", "Tìm kiếm sản phẩm", "Nhập tên hoặc thương hiệu", "Sản phẩm phù hợp được lọc", "Đạt"],
        ["5", "Xem chi tiết sản phẩm", "Chọn một sản phẩm", "Hiển thị ảnh, giá, thương hiệu, size, mô tả", "Đạt"],
        ["6", "Thêm vào giỏ hàng", "Chọn size và số lượng", "Sản phẩm được thêm vào giỏ theo đúng size", "Đạt"],
        ["7", "Đặt hàng", "Chọn sản phẩm, nhập địa chỉ", "Đơn hàng được tạo, sản phẩm đã đặt bị xóa khỏi giỏ", "Đạt"],
        ["8", "Quản lý sản phẩm", "Thêm/sửa/xóa sản phẩm", "Dữ liệu Product cập nhật đúng", "Đạt"],
        ["9", "Cập nhật đơn hàng", "Chuyển trạng thái theo luồng xử lý", "Trạng thái đơn thay đổi đúng thứ tự", "Đạt"],
        ["10", "Thống kê doanh số", "Mở trang thống kê hoặc lọc ngày", "Hiển thị doanh thu, top sản phẩm, trạng thái đơn", "Đạt"],
    ]
    add_table(doc, ["STT", "Chức năng", "Dữ liệu kiểm thử", "Kết quả mong đợi", "Trạng thái"], test_rows, widths_dxa=[700, 1900, 2700, 3000, 1060], font_size=10.3)

    add_heading(doc, "3.3.2 Đánh giá hệ thống", 3)
    add_para(doc, "Ưu điểm", bold=True)
    for item in [
        "Giao diện được chia rõ theo vai trò khách hàng, nhân viên và admin.",
        "Backend tách route và model tương đối rõ ràng, dễ theo dõi và mở rộng.",
        "Sản phẩm quản lý tồn kho theo từng size, phù hợp với đặc thù giày bóng đá.",
        "Luồng đơn hàng có quy tắc trạng thái rõ ràng và có kiểm tra thanh toán trước khi hoàn thành.",
        "Có thống kê doanh thu theo tháng, top sản phẩm bán chạy và trạng thái đơn hàng.",
    ]:
        add_bullet(doc, item)
    add_para(doc, "Hạn chế", bold=True)
    for item in [
        "Mật khẩu hiện lưu dạng chuỗi thường, chưa mã hóa bằng bcrypt.",
        "Chưa sử dụng JWT/session phía backend để bảo vệ API khi triển khai thật.",
        "Chưa tích hợp thanh toán trực tuyến và gửi email xác nhận đơn hàng.",
        "Chưa có chức năng đánh giá sản phẩm, mã giảm giá hoặc quản lý vận chuyển.",
        "Dữ liệu hình ảnh có thể lớn do lưu base64, cần cân nhắc lưu file hoặc dùng dịch vụ lưu trữ ảnh.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.3.3 Kết luận", 3)
    for text in [
        "Đề tài “Xây dựng website bán giày bóng đá TAF Soccer” đã vận dụng các kiến thức về HTML, CSS, JavaScript, Node.js, ExpressJS và MongoDB để tạo ra một hệ thống thương mại điện tử cơ bản nhưng có tính ứng dụng thực tế.",
        "Thông qua quá trình thực hiện, em hiểu rõ hơn cách tổ chức mã nguồn web, cách xây dựng REST API, cách thiết kế model dữ liệu bằng Mongoose, cách tích hợp frontend với backend và cách kiểm thử các chức năng chính của hệ thống.",
        "Trong tương lai, hệ thống có thể tiếp tục phát triển theo hướng bảo mật hơn, bổ sung thanh toán online, gửi email xác nhận, quản lý mã giảm giá, đánh giá sản phẩm và triển khai lên môi trường server thực tế.",
    ]:
        add_para(doc, text)

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Tài liệu MDN Web Docs về HTML, CSS và JavaScript.",
        "Tài liệu chính thức của Node.js và ExpressJS.",
        "Tài liệu MongoDB và Mongoose.",
        "Mã nguồn dự án TAF Soccer trong thư mục Backend và Frontend.",
        "Bài giảng môn JavaScript và Lập trình Web.",
    ]
    for ref in refs:
        add_number(doc, ref)


def build():
    diagrams = create_diagrams()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_toc(doc)
    add_intro(doc)
    add_chapter_1(doc)
    add_chapter_2(doc, diagrams)
    add_chapter_3(doc)
    doc.core_properties.title = "Báo cáo bài tập lớn - TAF Soccer"
    doc.core_properties.author = "Ngô Khuê Văn"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
